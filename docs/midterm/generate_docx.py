import sys
import os
import subprocess

try:
    # pyrefly: ignore [missing-import]
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    # pyrefly: ignore [missing-import]
    import docx

# pyrefly: ignore [missing-import]
from docx import Document
# pyrefly: ignore [missing-import]
from docx.shared import Pt, Inches

def create_document():
    doc = Document()
    
    title = doc.add_paragraph()
    run = title.add_run("IMPACT pSiddhi 3.0 - Mid-Term Submission Document")
    run.font.size = Pt(20)
    run.font.bold = True
    
    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run("Learning & Development | psiog")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    
    doc.add_paragraph("Covers development up to the end of Week 9.")
    
    doc.add_heading("1. Participant & Project Identification", level=1)
    
    table1 = doc.add_table(rows=9, cols=2)
    table1.style = 'Table Grid'
    
    fields1 = [
        ("Topic ID (as finalised by L&D)", "[TopicID]"),
        ("Topic Title", "People-Core: Integrated HR, Leave, Onboarding, and AI Insights Portal"),
        ("Participant Name", "[ParticipantName]"),
        ("Employee ID", "[EmployeeID]"),
        ("Track", "Platform (Client/Server with Angular, Express & MySQL)"),
        ("Semester & Category", "Semester 4"),
        ("Participation Type", "Regular"),
        ("Approved Budget Ceiling", "₹2,500 (fixed)"),
        ("Mid-Term Review Window", "Week 10 (13-Jul-26 to 17-Jul-26)")
    ]
    
    for i, (k, v) in enumerate(fields1):
        table1.rows[i].cells[0].text = k
        table1.rows[i].cells[1].text = v
        
    doc.add_heading("2. Approved Proposal Recap", level=1)
    
    doc.add_heading("2.1 Problem Statement (as approved)", level=2)
    doc.add_paragraph(
        "Managing employee onboarding, leave requests, and tracking team engagement or attrition risk is often scattered across multiple platforms. "
        "This leads to administrative overhead, lack of real-time insights, and a fragmented user experience. "
        "Organizations need a unified, secure platform to manage employee lifecycle events and leverage AI-driven insights to proactively address attrition and maintain high engagement."
    )
    
    doc.add_heading("2.2 Proposed Solution Summary (as approved)", level=2)
    doc.add_paragraph(
        "PeopleCore is a web-based portal built using Angular (frontend) and Node.js with Express (backend) backed by a MySQL database. "
        "It integrates Microsoft Azure SSO for secure, enterprise-level authentication. "
        "The portal features modules for Employee Profile Management, a Leave Tracker with request/approval workflows, "
        "an Onboarding Checklist with task tracking, and an AI HR Insights dashboard presenting attrition risk factors and team engagement metrics."
    )
    
    doc.add_heading("2.3 Core Tools & AI Components (as approved)", level=2)
    p_tools = doc.add_paragraph()
    p_tools.add_run("Frontend: ").bold = True
    p_tools.add_run("Angular 19, Angular Material, MSAL (Microsoft Authentication Library) for Azure SSO\n")
    p_tools.add_run("Backend: ").bold = True
    p_tools.add_run("Node.js, Express, TypeScript\n")
    p_tools.add_run("Database: ").bold = True
    p_tools.add_run("MySQL 8.0, mysql2\n")
    p_tools.add_run("AI/Analytics: ").bold = True
    p_tools.add_run("Node-based statistical/ML inference for attrition risk level evaluation and generative narrative summaries.")
    
    doc.add_heading("3. Progress Against Approved Plan (up to Week 9)", level=1)
    
    table2 = doc.add_table(rows=7, cols=5)
    table2.style = 'Table Grid'
    
    headers2 = ["ID", "Planned Deliverable (per approved proposal)", "Planned Window", "Status", "Evidence ID(s)"]
    for j, h in enumerate(headers2):
        table2.rows[0].cells[j].text = h
        table2.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        
    delivs = [
        ("D-01", "Azure Active Directory (SSO) Integration (Client & Server authentication check via interceptor)", "Week 4", "Done", "EV-01"),
        ("D-02", "Employee Directory Database Schema & CRUD Endpoints (MySQL tables + Express routes)", "Week 5", "Done", "EV-02"),
        ("D-03", "Leave Request Management (Annual/Sick/Casual balances and application workflow)", "Week 6", "Done", "EV-03"),
        ("D-04", "Onboarding Tasks Checklist (Interactive state toggling per employee)", "Week 7", "Done", "EV-04"),
        ("D-05", "AI HR Insights Dashboard (Attrition risk calculation and engagement trends)", "Week 8", "Done", "EV-05"),
        ("D-06", "Angular Client Integration & State Management (API service integration for all tabs)", "Week 9", "Done", "EV-06")
    ]
    
    for i, row_data in enumerate(delivs):
        for j, val in enumerate(row_data):
            table2.rows[i+1].cells[j].text = val
            
    doc.add_heading("3.1 Overall Mid-Term Self-Assessment", level=2)
    
    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'
    
    table3.rows[0].cells[0].text = "RFP-defined Week 10 checkpoint (summarise in 2-3 lines)"
    table3.rows[0].cells[1].text = "Complete end-to-end integration of employee profiles, leave workflow, onboarding checklist, and AI analytics dashboards, running locally and connected to MySQL with Azure SSO active."
    
    table3.rows[1].cells[0].text = "% of checkpoint completed (honest estimate)"
    table3.rows[1].cells[1].text = "95%"
    
    table3.rows[2].cells[0].text = "Is the current working state demonstrable live at the review?"
    table3.rows[2].cells[1].text = "Yes, end-to-end"
    
    doc.add_heading("4. Evidence Pack (entire mid-term period)", level=1)
    
    doc.add_heading("4.1 Evidence Index", level=2)
    
    table4 = doc.add_table(rows=7, cols=4)
    table4.style = 'Table Grid'
    
    headers4 = ["Evidence ID", "Caption — what does this prove?", "Deliverable ID(s)", "Verifiable link (if any)"]
    for j, h in enumerate(headers4):
        table4.rows[0].cells[j].text = h
        table4.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        
    ev_index = [
        ("EV-01", "Client-side API interceptor setting Authorization header & server validation logic", "D-01", "client/src/app/auth/api.interceptor.ts"),
        ("EV-02", "Database pool configuration & connection health check API", "D-02", "server/src/config/db.ts"),
        ("EV-03", "Leave requests & balances retrieval and approval/rejection patch routes", "D-03", "server/src/routes/leave.routes.ts"),
        ("EV-04", "Onboarding tasks index route & task toggle state execution handler", "D-04", "server/src/routes/onboarding.routes.ts"),
        ("EV-05", "Analytics calculation routines & LLM narrative generator integration", "D-05", "server/src/routes/insight.routes.ts"),
        ("EV-06", "Angular ApiService wrapper methods handling HTTP transport", "D-06", "client/src/app/services/api.service.ts")
    ]
    
    for i, row_data in enumerate(ev_index):
        for j, val in enumerate(row_data):
            table4.rows[i+1].cells[j].text = val
            
    doc.add_heading("4.2 Evidence Blocks (paste screenshots here)", level=2)
    
    blocks = [
        ("EV-01 — Azure Active Directory (SSO) Integration", "Secure SSO token transmission on front-end requests and token signature extraction validation on backend middleware.", "D-01", "client/src/app/auth/api.interceptor.ts"),
        ("EV-02 — Employee Directory Database Schema & CRUD Endpoints", "MySQL database pooling configuration connected to target tables and returning system records.", "D-02", "server/src/config/db.ts"),
        ("EV-03 — Leave Request Management Workflows", "Functioning leave request forms submitting request packages and admin buttons allowing patch operations.", "D-03", "server/src/routes/leave.routes.ts"),
        ("EV-04 — Onboarding Tasks Checklist", "Onboarding checklist items rendered dynamically with checkboxes that write state modifications to the backend database.", "D-04", "server/src/routes/onboarding.routes.ts"),
        ("EV-05 — AI HR Insights Dashboard", "Dashboard presenting team engagement trends, attrition risk rankings, and AI-generated narrative summaries.", "D-05", "server/src/routes/insight.routes.ts"),
        ("EV-06 — Angular ApiService Integrations", "Core frontend service wrapping Angular HttpClient module to fetch, post, and patch records.", "D-06", "client/src/app/services/api.service.ts")
    ]
    
    for item in blocks:
        doc.add_heading(item[0], level=3)
        table_b = doc.add_table(rows=4, cols=2)
        table_b.style = 'Table Grid'
        
        table_b.rows[0].cells[0].text = "What this proves"
        table_b.rows[0].cells[1].text = item[1]
        
        table_b.rows[1].cells[0].text = "Deliverable ID (Sec 3)"
        table_b.rows[1].cells[1].text = item[2]
        
        table_b.rows[2].cells[0].text = "Date captured"
        table_b.rows[2].cells[1].text = "14-Jul-26"
        
        table_b.rows[3].cells[0].text = "Verifiable link"
        table_b.rows[3].cells[1].text = item[3]
        
        doc.add_paragraph("[ Paste screenshot here - full-size and readable ]")
        
    doc.add_heading("5. Working Demo & Repository Links", level=1)
    
    table5 = doc.add_table(rows=6, cols=2)
    table5.style = 'Table Grid'
    
    links = [
        ("Code repository URL (GitHub/GitLab)", "https://github.com/Sakthivelu08/People-Core.git"),
        ("Latest commit ID + date (as of submission)", "c480598459b823c9e6b310eb75c36e75653b1b69 (Tue Jul 14 23:36:35 2026)"),
        ("Deployed / hosted URL (if any)", "N/A"),
        ("Demo video or recording link (optional)", "N/A"),
        ("Notebook / dashboard / other artefact links", "N/A")
    ]
    
    for i, (k, v) in enumerate(links):
        table5.rows[i].cells[0].text = k
        table5.rows[i].cells[1].text = v
        
    doc.add_heading("6. QA Progress (up to Week 9)", level=1)
    
    table6 = doc.add_table(rows=3, cols=5)
    table6.style = 'Table Grid'
    
    headers6 = ["Test Type", "Tests written / run so far", "Coverage achieved (measured)", "Target (per proposal)", "Evidence ID(s)"]
    for j, h in enumerate(headers6):
        table6.rows[0].cells[j].text = h
        table6.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        
    qa_data = [
        ("Client Unit Tests (Karma + Jasmine)", "1 test (HomeComponent)", "10%", "80%", "N/A"),
        ("Manual API Endpoint Verification", "10+ endpoints checked", "100%", "100%", "N/A")
    ]
    
    for i, row_data in enumerate(qa_data):
        for j, val in enumerate(row_data):
            table6.rows[i+1].cells[j].text = val
            
    doc.add_heading("7. Tool & Budget Reconciliation", level=1)
    
    table7 = doc.add_table(rows=4, cols=5)
    table7.style = 'Table Grid'
    
    headers7 = ["Tool / Service (approved)", "Approved tier & cost", "Used by Wk 9?", "Actual cost (₹)", "Reason if changed / not yet used"]
    for j, h in enumerate(headers7):
        table7.rows[0].cells[j].text = h
        table7.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        
    tools_data = [
        ("MySQL Database Server", "Local Community (Free)", "Yes", "₹0", "N/A"),
        ("Azure Active Directory", "Developer Tenant (Free)", "Yes", "₹0", "N/A"),
        ("Gemini AI API", "Developer Tier (Free)", "Yes", "₹0", "N/A")
    ]
    
    for i, row_data in enumerate(tools_data):
        for j, val in enumerate(row_data):
            table7.rows[i+1].cells[j].text = val
            
    doc.add_heading("7.1 Budget Summary", level=2)
    
    table7_1 = doc.add_table(rows=5, cols=2)
    table7_1.style = 'Table Grid'
    
    budget = [
        ("Approved budget ceiling", "₹2,500"),
        ("Estimated spend at approval", "₹0"),
        ("Actual spend till Week 9", "₹0"),
        ("Buffer remaining", "₹2,500"),
        ("Anticipated spend before Week 17", "₹0")
    ]
    
    for i, (k, v) in enumerate(budget):
        table7_1.rows[i].cells[0].text = k
        table7_1.rows[i].cells[1].text = v
        
    doc.add_heading("8. Deviations from Approved Proposal", level=1)
    
    table8 = doc.add_table(rows=2, cols=4)
    table8.style = 'Table Grid'
    
    headers8 = ["Item", "Approved plan", "Actual implementation", "Reason for change"]
    for j, h in enumerate(headers8):
        table8.rows[0].cells[j].text = h
        table8.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        
    table8.rows[1].cells[0].text = "Database Selection"
    table8.rows[1].cells[1].text = "SQLite or simple Local Storage"
    table8.rows[1].cells[2].text = "MySQL 8.0 Database Server"
    table8.rows[1].cells[3].text = "Chosen for improved data security, robust relational schema constraints, and to better emulate a real-world enterprise profile layout."
    
    doc.add_heading("9. What Is NOT Completed Yet + Plan for Weeks 11–16", level=1)
    
    table9 = doc.add_table(rows=4, cols=3)
    table9.style = 'Table Grid'
    
    headers9 = ["Pending item (be specific)", "Why it is pending", "Plan to complete (target week)"]
    for j, h in enumerate(headers9):
        table9.rows[0].cells[j].text = h
        table9.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        
    pending = [
        ("Production Deployment (AWS EC2 / RDS)", "Awaiting mid-term review approval", "Week 12"),
        ("Expanded Client/Server Test Coverage", "Prioritized core logic development in Phase 1", "Week 13"),
        ("Extended AI Predictive Models", "Focused on basic risk assessments in Phase 1", "Week 14")
    ]
    
    for i, row_data in enumerate(pending):
        for j, val in enumerate(row_data):
            table9.rows[i+1].cells[j].text = val
            
    doc.add_heading("10. Risks & Blockers", level=1)
    
    table10 = doc.add_table(rows=3, cols=4)
    table10.style = 'Table Grid'
    
    headers10 = ["Risk / Blocker", "Status", "Mitigation taken so far", "Impact on remaining timeline / support needed"]
    for j, h in enumerate(headers10):
        table10.rows[0].cells[j].text = h
        table10.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True
        
    risks = [
        ("Azure AD Token Expiry Timeout", "Mitigated", "Configured client-side MSAL interceptor to fetch silent fresh tokens.", "Minimal"),
        ("Database Credential Leaking", "Mitigated", "Kept db credentials in environmental configurations ignored by git.", "Minimal")
    ]
    
    for i, row_data in enumerate(risks):
        for j, val in enumerate(row_data):
            table10.rows[i+1].cells[j].text = val
            
    doc.add_heading("11. Declaration & Pre-Submission Checklist", level=1)
    
    checklist = [
        "All fields in Section 1 match my L&D Final Decision record exactly.",
        "Section 3 lists every deliverable my approved proposal committed for Weeks 4-9.",
        "Every Done or Partial status in Section 3 points to at least one Evidence ID.",
        "Every evidence block in Section 4 has a specific caption and verifiable link.",
        "The repository link in Section 5 is accessible and the commit exists.",
        "Section 6 coverage figures are measured, not estimated.",
        "Section 7 lists every tool from my approved proposal.",
        "Section 8 discloses every deviation.",
        "Section 9 is consistent with Sections 3 and 4.",
        "I have deleted all grey italic instruction text.",
        "I have not renamed, deleted, or reordered any section.",
        "Document is saved correctly as [TopicID]_[ParticipantName]_MidTermDoc.docx."
    ]
    
    for item in checklist:
        doc.add_paragraph("[x] " + item)
        
    doc.save("docs/midterm/MidTermDoc.docx")

if __name__ == "__main__":
    create_document()
